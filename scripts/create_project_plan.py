from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/Rahman-Plastic-Surgery-Website-Implementation-Plan.docx")

NAVY = "2B231E"
BLUE = "A88D70"
INK = "2C2521"
BODY = "665D57"
MUTED = "9F8F80"
LINE = "E5DCD2"
SOFT = "F1E7DD"
PALE = "F8F2EB"
WHITE = "FFFDFB"
GREEN = "456E57"


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:{}".format(key)), str(value))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_run(run, **kwargs)
    return run


def add_paragraph(doc, text="", style=None, before=0, after=6, line=1.25, color=INK):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if text:
        add_text(p, text, size=11, color=color)
    return p


def add_bullets(doc, items):
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        add_text(p, text, size=10.5, color=INK)


def add_numbered(doc, items):
    for text in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        add_text(p, text, size=10.5, color=INK)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_text(p, text, size={1: 16, 2: 13, 3: 11.5}[level], color=NAVY if level == 1 else BLUE, bold=True)
    return p


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, SOFT)
    set_cell_border(cell, top={"val": "single", "sz": "6", "color": BLUE}, left={"val": "single", "sz": "6", "color": BLUE}, bottom={"val": "single", "sz": "6", "color": BLUE}, right={"val": "single", "sz": "6", "color": BLUE})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_text(p, title + "  ", size=10.5, color=NAVY, bold=True)
    add_text(p, text, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    header = table.rows[0]
    for i, label in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, top={"val": "single", "sz": "4", "color": NAVY}, left={"val": "single", "sz": "4", "color": NAVY}, bottom={"val": "single", "sz": "4", "color": NAVY}, right={"val": "single", "sz": "4", "color": NAVY})
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_text(p, label, size=9.5, color=WHITE, bold=True)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cell = cells[i]
            set_cell_shading(cell, WHITE if row_idx % 2 == 0 else PALE)
            set_cell_border(cell, top={"val": "single", "sz": "4", "color": LINE}, left={"val": "single", "sz": "4", "color": LINE}, bottom={"val": "single", "sz": "4", "color": LINE}, right={"val": "single", "sz": "4", "color": LINE})
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_text(p, value, size=9.25, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run(run, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, color, before, after in ((1, 16, NAVY, 16, 8), (2, 13, BLUE, 12, 6), (3, 11.5, NAVY, 8, 4)):
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.3)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    add_text(p, "RAHMAN PLASTIC SURGERY  /  WEBSITE IMPLEMENTATION PLAN", size=8.5, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(fp, "CONFIDENTIAL PROJECT PLAN  •  ", size=8.5, color=MUTED)
    add_page_number(fp)


def build_document():
    doc = Document()
    configure_document(doc)

    # Customer-pack opening pattern: left-aligned title and compact metadata grid.
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(28)
    kicker.paragraph_format.space_after = Pt(8)
    add_text(kicker, "PROJECT DELIVERY PLAN", size=10, color=BLUE, bold=True)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    add_text(title, "Rahman Plastic Surgery", size=28, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(20)
    add_text(subtitle, "Premium medical tourism website — implementation plan for version one", size=14, color=BODY)

    meta = doc.add_table(rows=2, cols=2)
    set_table_geometry(meta, [1800, 7560])
    for row, (label, value) in zip(meta.rows, [("Platform", "Next.js 16 • JavaScript • Tailwind CSS 4"), ("Prepared", date.today().strftime("%d %B %Y"))]):
        set_cell_shading(row.cells[0], SOFT)
        set_cell_shading(row.cells[1], WHITE)
        for cell in row.cells:
            set_cell_border(cell, top={"val": "single", "sz": "4", "color": LINE}, left={"val": "single", "sz": "4", "color": LINE}, bottom={"val": "single", "sz": "4", "color": LINE}, right={"val": "single", "sz": "4", "color": LINE})
        add_text(row.cells[0].paragraphs[0], label, size=9.5, color=NAVY, bold=True)
        add_text(row.cells[1].paragraphs[0], value, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_callout(doc, "Plan intent", "Deliver a premium, responsive, conversion-focused single-page clinic website that serves patients in Pakistan and international visitors, while keeping all unverified medical information and patient content clearly marked as provisional.")

    add_heading(doc, "1. Delivery outcome", 1)
    add_paragraph(doc, "The version-one site will present Rahman Plastic Surgery as a calm, modern, internationally credible private clinic. It will guide visitors from procedure discovery through consultation intent, with phone and WhatsApp routes always available as low-friction alternatives.")
    add_bullets(doc, [
        "Single-page App Router experience with anchor navigation, mobile-first layouts, and centralised editable content.",
        "Premium editorial design: deep cocoa for trust-building sections, muted taupe for calls to action, and generous warm-white space.",
        "A validated consultation form with an API endpoint that only sends email once approved Resend credentials are available.",
        "Accessible interaction patterns: keyboard-operable menu, slider, and modal; visible focus; reduced-motion support; and 44px touch targets.",
        "No invented credentials, prices, reviews, statistics, or patient outcomes. Placeholder material remains visibly labelled until approved assets arrive.",
    ])

    add_heading(doc, "Reference-led visual direction", 2)
    add_paragraph(doc, "The supplied clinic references set the visual standard for version one. They are direction, not a template to copy: a composed, image-led, editorial medical experience with human warmth and restrained clinical confidence.")
    add_table(doc, ["Design element", "Confirmed direction"], [
        ("Palette", "Warm white and paper surfaces, deep cocoa sections, muted taupe as the only accent, and no electric blue, purple, or decorative gradients."),
        ("Typography", "DM Sans for interface copy and Instrument Sans for display headings. Keep hierarchy calm, with weight and spacing rather than oversized type."),
        ("Section rhythm", "Use white/paper editorial split sections with full-height clinical imagery. Alternate sparingly with deep cocoa narrative sections. Use slim tabs and quiet rules instead of card walls."),
        ("Calls to action", "Use compact rectangular taupe or paper buttons with uppercase labels. Reserve pill shapes for compact icon controls only."),
        ("Media & proof", "Show only approved clinic, team, travel, and consented patient materials. Do not use stock or AI imagery as proof of surgical outcomes."),
    ], [2200, 7160])
    add_callout(doc, "Anti-template standard", "Avoid generic three-card layouts, glass panels, neon glow, decorative motion, fake testimonials, visible development labels, numbered eyebrows, and AI-style filler language. The site should feel intentional, quiet, and specific to a private surgical practice.")

    add_heading(doc, "2. Scope and information architecture", 1)
    add_paragraph(doc, "The homepage follows the supplied conversion order. Each section has a distinct job, and shared content stays in data modules rather than being repeated in page markup.")
    add_table(doc, ["Experience stage", "Homepage sections", "Visitor outcome"], [
        ("Trust & orientation", "Hero; trust strip; why choose us", "Understand the clinic's position, care approach, and next action."),
        ("Procedure discovery", "Featured procedures; before & after", "Explore body-contouring options and safe, consent-led outcomes context."),
        ("International conversion", "International patients; 5-step journey; patient services; why Pakistan", "Understand travel coordination, consultation route, and care planning."),
        ("Proof & reassurance", "Testimonials; strong CTA; recovery journey; patient journey gallery; statistics", "Build confidence without unsupported proof points or claims."),
        ("Contact & handoff", "Location/map; consultation form; footer", "Choose call, WhatsApp, or a validated consultation request."),
    ], [1750, 4100, 3510])

    add_heading(doc, "3. Technical approach", 1)
    add_table(doc, ["Area", "Implementation decision"], [
        ("Foundation", "Retain the existing Next.js 16 App Router project, JavaScript, and Tailwind CSS 4. No TypeScript conversion or unnecessary configuration replacement."),
        ("Components", "Create reusable layout, section, and UI components; keep page.js as a composition layer. Data files hold procedures, services, testimonials, stats, gallery entries, and placeholder labels."),
        ("Packages", "Add lucide-react, react-hook-form, and zod only if absent. Use CSS transitions, native scroll snapping, and a custom range slider—no heavy carousel or animation library."),
        ("Contact workflow", "Use site.js helpers for phone, email, and WhatsApp URLs. API route validates the same Zod schema server-side, rejects honeypots, and conditionally sends email through Resend."),
        ("Images", "Use next/image throughout; hero alone is priority-loaded. Development imagery lives in public/images and all replace-before-launch assets are labelled in the relevant data files."),
        ("SEO", "Add page metadata, Open Graph placeholders, and a minimal JSON-LD helper structure that excludes unverified business details, reviews, prices, and credentials."),
    ], [2000, 7360])

    add_heading(doc, "4. Proposed project structure", 1)
    add_paragraph(doc, "The exact directories will follow the current repository where practical; the target structure below keeps presentation, content, and integration concerns separate.")
    structure = [
        "src/app/ — layout, homepage composition, globals, and api/consultation/route.js",
        "src/components/layout/ — announcement bar, header, mobile menu, mobile conversion bar, footer, and back-to-top",
        "src/components/sections/ — one component per required content section",
        "src/components/ui/ — Button, SectionHeading, cards, modal, image-label, and BeforeAfterSlider",
        "src/data/ — site config plus procedures, services, testimonials, comparisons, gallery, and statistics data",
        "src/lib/ — URL helpers, validation schema, and low-level utilities",
        "public/images/ — development placeholders grouped by hero, procedures, before-after, journey, team, and testimonials",
    ]
    add_bullets(doc, structure)

    add_heading(doc, "5. Delivery roadmap", 1)
    add_paragraph(doc, "The plan is organised in implementation order. Each phase leaves the site in a testable state and avoids waiting on client assets where a clearly labelled placeholder can be used.")
    add_table(doc, ["Phase", "Work", "Exit criteria"], [
        ("01\nFoundation", "Audit; add approved missing packages; create site config, tokens, URL helpers, and validation schema.", "Clean build; central contact helpers and tokens available."),
        ("02\nShell", "Build navigation, global CTAs, mobile menu, sticky conversion bar, footer, and anchor routes.", "Every CTA resolves; keyboard menu works with Escape."),
        ("03\nCore story", "Build hero, trust, procedures, international care, journey, services, and comparison sections.", "Responsive narrative complete with safe editable copy."),
        ("04\nConfidence", "Build before/after slider, demo testimonials, CTA, recovery, gallery/lightbox, and stats.", "No placeholder can be mistaken for verified evidence."),
        ("05\nLead capture", "Build location placeholder, form, client/server validation, honeypot, and conditional Resend path.", "Clear form states; no false email-success reporting."),
        ("06\nQA & handoff", "Add metadata, test responsive and keyboard paths, resolve warnings, write README and launch checklist.", "No overflow, broken images, or unresolved core accessibility issues."),
    ], [1450, 5000, 2910])

    add_heading(doc, "6. Content, safety, and compliance guardrails", 1)
    add_callout(doc, "Non-negotiable", "The site must never turn illustration into evidence. Until client approval is provided, patient images, testimonials, doctor credentials, statistics, costs, and results must stay editable and visibly identified as illustrative, demo, or to be confirmed.")
    add_bullets(doc, [
        "Use a contextual medical disclaimer in the footer. Do not imply that web content replaces professional consultation or guarantees a result.",
        "Before/after layouts display only consented patient materials when supplied; stock or placeholder assets receive a clear illustrative label.",
        "Testimonials use the provided placeholder wording until written approval and review source details are received.",
        "Statistics are data-driven and must be verified by the client before launch; do not transfer claims from visual references.",
        "Do not add patient photo uploads in version one. That requires secure storage, consent, retention, and privacy processes.",
    ])

    add_heading(doc, "7. Consultation request workflow", 1)
    add_numbered(doc, [
        "Visitor completes required personal details, procedure interest, international-patient status, optional message, and privacy consent.",
        "Client-side Zod validation gives accessible field-level feedback before submission.",
        "POST /api/consultation validates the same data server-side and rejects honeypot entries.",
        "When RESEND_API_KEY, CONTACT_EMAIL, and FROM_EMAIL are configured, the route sends the request to the clinic contact inbox.",
        "Without valid production credentials, the route returns an honest error; the UI directs the visitor to WhatsApp or phone instead of claiming a request was sent.",
    ])

    add_heading(doc, "8. Client inputs required before launch", 1)
    add_table(doc, ["Input", "Why it matters", "Launch status"], [
        ("Logo (SVG or transparent PNG)", "Replaces temporary RPS monogram/text mark.", "Required"),
        ("Exact phone, WhatsApp, email, address, hours, maps URL", "Activates accurate contact paths, location card, and map embed.", "Required"),
        ("Surgeon credentials, team details, clinic imagery", "Supports trustworthy, approved clinical representation.", "Required for final content"),
        ("Consented before/after photos and patient journey photos", "Replaces illustrative layouts without misrepresenting outcomes.", "Required for outcomes gallery"),
        ("Approved testimonials/video reviews and Google Reviews URL", "Replaces demo testimonial content.", "Required for social proof"),
        ("Procedure copy, recovery guidance, and optional verified cost ranges", "Confirms medically appropriate procedure detail and comparison context.", "Required for final clinical content"),
        ("Domain and Resend credentials", "Enables canonical URL, production email, and final sender configuration.", "Required for live lead delivery"),
    ], [2700, 4980, 1680])

    add_heading(doc, "9. Open decisions to resolve", 1)
    add_table(doc, ["Decision", "Current state", "Recommended resolution"], [
        ("Typography", "The active project direction uses DM Sans + Instrument Sans, chosen to match the approved visual references.", "Keep this pair as the single type system unless the client supplies a formal brand font."),
        ("Verified claims", "Experience, surgeon-count, facility, and international support claims are supplied as editable draft content.", "Approve exact wording and supporting evidence before publishing."),
        ("Consultation email", "No credentials are currently supplied.", "Provide Resend API key, verified sender, and destination inbox; complete an end-to-end delivery test."),
        ("Maps and legal pages", "No final clinic location, Privacy Policy, Terms, or Medical Disclaimer URLs are provided.", "Supply approved URLs/content before production release."),
    ], [1900, 3750, 3710])

    add_heading(doc, "10. Quality gates and acceptance checklist", 1)
    add_bullets(doc, [
        "Run npm run lint and npm run build with no errors; resolve browser console, hydration, and Next Image warnings.",
        "Review at 1440px, 1280px, 1024px, 768px, 430px, 390px, 375px, and 360px; confirm no horizontal overflow.",
        "Confirm header switches below 1024px, procedure cards switch below 768px, hero CTAs stack below 430px, and the sticky mobile bar never hides form actions.",
        "Verify every CTA route, central WhatsApp helper, phone link, mailto link, focus state, Escape interaction, and keyboard path.",
        "Test form success/error states, server-side validation, honeypot rejection, and missing-credential behaviour.",
        "Confirm that all placeholder content is visibly labelled and documented in the README, alongside contact/content locations, environment variables, and pre-launch tasks.",
    ])

    add_heading(doc, "11. Definition of ready for launch", 1)
    add_paragraph(doc, "The site is ready for production only when the technical quality gates pass and the client has supplied or approved all public-facing medical, legal, contact, image, testimonial, and map content. The build can progress with placeholders, but launch should wait for the approvals listed in Section 8.")

    final_note = doc.add_paragraph()
    final_note.paragraph_format.space_before = Pt(16)
    final_note.paragraph_format.space_after = Pt(0)
    add_text(final_note, "Prepared from the supplied project brief. This plan is an implementation guide, not a source of medical claims or patient-facing clinical content.", size=9, color=MUTED, italic=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_document()
